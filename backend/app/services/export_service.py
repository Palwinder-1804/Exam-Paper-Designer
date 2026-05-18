import json
import os
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors

from docx import Document
from docx.shared import Inches, Pt

from typing import Any, Dict, List, Optional, Union

FIG_DIR = "app/static/figures/current"
FIGURES_JSON_PATH = os.path.join(FIG_DIR, "figures.json")


def _load_figures_map() -> Dict[str, str]:
    if not os.path.exists(FIGURES_JSON_PATH):
        return {}
    try:
        with open(FIGURES_JSON_PATH, "r", encoding="utf-8") as f:
            data = json.load(f)
        if not isinstance(data, list):
            return {}
        out: Dict[str, str] = {}
        for item in data:
            if isinstance(item, dict) and item.get("id") and item.get("filename"):
                out[item["id"]] = item["filename"]
        return out
    except Exception:
        return {}


def _render_reportlab_image(
    img_path: str,
    max_width: float = 4.2 * inch,
    max_height: float = 3.2 * inch,
) -> RLImage:
    """Scale image to fit exam-paper column without overflow."""
    try:
        from PIL import Image as PILImage

        pil = PILImage.open(img_path)
        w_px, h_px = pil.size
        if w_px and h_px:
            aspect = h_px / float(w_px)
            w = max_width
            h = w * aspect
            if h > max_height:
                h = max_height
                w = h / aspect
            return RLImage(img_path, width=w, height=h)
    except Exception:
        pass
    return RLImage(img_path, width=max_width, height=max_height)


def export_pdf(content: Union[str, List[Dict[str, Any]]], filename: str = "paper.pdf"):
    folder = "app/static/outputs"
    os.makedirs(folder, exist_ok=True)

    path = f"{folder}/{filename}"

    doc = SimpleDocTemplate(path)
    styles = getSampleStyleSheet()
    header_style = ParagraphStyle(
        "HeaderStyle",
        parent=styles["Normal"],
        fontSize=14,
        leading=18,
        spaceAfter=12,
        alignment=1,
    )
    section_style = ParagraphStyle(
        "SectionStyle",
        parent=styles["Normal"],
        fontSize=12,
        leading=16,
        spaceBefore=10,
        spaceAfter=8,
    )
    q_title_style = ParagraphStyle(
        "QuestionTitleStyle",
        parent=styles["Normal"],
        fontSize=11,
        leading=14,
        spaceAfter=6,
    )
    q_text_style = ParagraphStyle(
        "QuestionTextStyle",
        parent=styles["Normal"],
        fontSize=11,
        leading=14,
        leftIndent=20,
        spaceAfter=6,
    )

    elements = []

    figures_map = _load_figures_map()

    inst_style = ParagraphStyle(
        "InstitutionStyle",
        parent=styles["Normal"],
        fontSize=13,
        leading=16,
        alignment=1,
        spaceAfter=6,
        textColor=colors.HexColor("#1a365d"),
    )

    if isinstance(content, str):
        for line in content.split("\n"):
            if line.strip():
                elements.append(Paragraph(line, styles["Normal"]))
                elements.append(Spacer(1, 10))
    else:
        for item in content:
            item_type = item.get("type")

            if item_type == "institution":
                elements.append(Paragraph(item.get("text", ""), inst_style))
                elements.append(Spacer(1, 4))
                continue

            if item_type in ("subject_line", "student_fields"):
                elements.append(Paragraph(item.get("text", ""), q_title_style))
                elements.append(Spacer(1, 6))
                continue

            if item_type == "divider":
                elements.append(Spacer(1, 8))
                continue

            if item_type == "section_instruction":
                si_style = ParagraphStyle(
                    "SectionInstr",
                    parent=styles["Italic"],
                    fontSize=10,
                    leading=13,
                    leftIndent=12,
                    spaceAfter=8,
                    textColor=colors.dimgrey,
                )
                elements.append(Paragraph(item.get("text", "").replace("\n", "<br/>"), si_style))
                continue
            
            if item_type == "header_title":
                elements.append(Paragraph(item.get("text", ""), header_style))
                elements.append(Spacer(1, 8))
                continue
                
            if item_type == "header_meta":
                elements.append(Paragraph(item.get("text", ""), q_title_style))
                elements.append(Spacer(1, 8))
                continue
                
            if item_type == "instructions":
                instructions_style = ParagraphStyle(
                    "InstructionStyle",
                    parent=styles["Italic"],
                    fontSize=10,
                    leading=14,
                    spaceAfter=12,
                    textColor=colors.dimgrey
                )
                txt = item.get("text", "").replace("\n", "<br/>")
                elements.append(Paragraph(txt, instructions_style))
                continue

            if item_type == "section":
                elements.append(Paragraph(item.get("text", ""), section_style))
                elements.append(Spacer(1, 4))
                continue

            if item_type == "question":
                qn = item.get("number")
                marks = item.get("marks")
                elements.append(Paragraph(f"Q{qn}. ({marks} marks)", q_title_style))

                for b in item.get("blocks", []) or []:
                    b_type = b.get("type")
                    if b_type == "text":
                        txt = (b.get("text") or "").strip()
                        if txt:
                            elements.append(Paragraph(txt.replace("\n", "<br/>"), q_text_style))
                    elif b_type == "sub_questions":
                        for sq in b.get("items") or []:
                            label = sq.get("label", "a")
                            stxt = (sq.get("text") or "").strip()
                            if stxt:
                                elements.append(
                                    Paragraph(
                                        f"({label}) {stxt.replace(chr(10), '<br/>')}",
                                        q_text_style,
                                    )
                                )
                    elif b_type == "options":
                        for i, opt in enumerate(b.get("options") or []):
                            label = chr(ord("a") + i)
                            elements.append(
                                Paragraph(f"({label}) {opt}", q_text_style)
                            )
                    elif b_type == "figure":
                        fig_id = (b.get("figure_id") or b.get("id") or "").strip()
                        caption = (b.get("caption") or "Figure").strip()
                        filename_for_fig = figures_map.get(fig_id)
                        if filename_for_fig:
                            img_path = os.path.join(FIG_DIR, filename_for_fig)
                            if os.path.exists(img_path):
                                elements.append(Paragraph(f"<i>{caption}</i>", q_text_style))
                                elements.append(Spacer(1, 4))
                                elements.append(_render_reportlab_image(img_path))
                                elements.append(Spacer(1, 10))
                            else:
                                elements.append(
                                    Paragraph(
                                        f"[Diagram unavailable: {fig_id}]",
                                        q_text_style,
                                    )
                                )

    doc.build(elements)

    return path

def export_docx(content: Union[str, List[Dict[str, Any]]], filename: str = "paper.docx"):
    folder = "app/static/outputs"
    os.makedirs(folder, exist_ok=True)

    path = f"{folder}/{filename}"

    doc = Document()
    figures_map = _load_figures_map()

    if isinstance(content, str):
        for line in content.split("\n"):
            doc.add_paragraph(line)
    else:
        for item in content:
            item_type = item.get("type")

            if item_type == "institution":
                p = doc.add_paragraph(item.get("text", ""))
                p.runs[0].bold = True
                p.alignment = 1
                continue

            if item_type in ("subject_line", "student_fields", "divider"):
                doc.add_paragraph(item.get("text", ""))
                continue

            if item_type == "section_instruction":
                p = doc.add_paragraph(item.get("text", ""))
                p.runs[0].italic = True
                continue
            
            if item_type == "header_title":
                p = doc.add_paragraph(item.get("text", ""))
                p.runs[0].bold = True
                p.paragraph_format.space_after = Pt(12)
                p.alignment = 1 # Center
                
            elif item_type == "header_meta":
                p = doc.add_paragraph(item.get("text", ""))
                p.runs[0].bold = True
                p.paragraph_format.space_after = Pt(12)
                p.alignment = 1 # Center
                
            elif item_type == "instructions":
                p = doc.add_paragraph(item.get("text", ""))
                p.runs[0].italic = True
                p.paragraph_format.space_after = Pt(12)
                
            elif item_type == "section":
                ps = doc.add_paragraph(item.get("text", ""))
                ps.runs[0].bold = True
                ps.paragraph_format.space_after = Pt(6)
            elif item_type == "question":
                qn = item.get("number")
                marks = item.get("marks")
                pq = doc.add_paragraph(f"Q{qn}. ({marks} marks)")
                pq.runs[0].bold = True
                pq.paragraph_format.space_after = Pt(4)

                for b in item.get("blocks", []) or []:
                    b_type = b.get("type")
                    if b_type == "text":
                        txt = (b.get("text") or "").strip()
                        if txt:
                            pt = doc.add_paragraph(txt)
                            pt.paragraph_format.left_indent = Inches(0.25)
                    elif b_type == "sub_questions":
                        for sq in b.get("items") or []:
                            label = sq.get("label", "a")
                            stxt = (sq.get("text") or "").strip()
                            if stxt:
                                psq = doc.add_paragraph(f"({label}) {stxt}")
                                psq.paragraph_format.left_indent = Inches(0.35)
                    elif b_type == "options":
                        for i, opt in enumerate(b.get("options") or []):
                            label = chr(ord("a") + i)
                            po = doc.add_paragraph(f"({label}) {opt}")
                            po.paragraph_format.left_indent = Inches(0.35)
                    elif b_type == "figure":
                        fig_id = (b.get("figure_id") or b.get("id") or "").strip()
                        caption = (b.get("caption") or "Figure").strip()
                        filename_for_fig = figures_map.get(fig_id)
                        if filename_for_fig:
                            img_path = os.path.join(FIG_DIR, filename_for_fig)
                            if os.path.exists(img_path):
                                pc = doc.add_paragraph(caption)
                                pc.paragraph_format.left_indent = Inches(0.25)
                                doc.add_picture(img_path, width=Inches(4.2))
                            else:
                                doc.add_paragraph(f"[Diagram unavailable: {fig_id}]")

    doc.save(path)

    return path